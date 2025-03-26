import re
import os
import time
import warnings
import dateparser
import glob
from datetime import datetime
import polars as pl
import numpy as np
from docx import Document, text
from nlp_funcs import find_persons
from custom_database import CsvDatabase

def find_study_pattern(text):
    pattern = r"[^{(_\s.]{2,6}_\d{2,3}_\d{1,2}[A-Za-z]*\d{2,4}"
    pat = re.search(pattern, text)

    return pat

def parse_custom_date(date_str):
    # Mapping of month abbreviations to their numeric representation
    month_map = {
        'JAN': 1, 'FEB': 2, 'MAR': 3, 'MARCH': 3, 'APR': 4, 'APL': 4, 'ARP': 4,
        'APRIL': 4, 'MAY': 5, 'JUN': 6, 'JUNE': 6, 'JUL': 7, 'JULY': 7, 'AUG': 8,
        'SEP': 9, 'SEPT': 9, 'OCT': 10, 'NOV': 11, 'DEC': 12,
    }

    nums = re.findall('\d+', date_str)
    month = re.findall('[A-Za-z]+', date_str)

    # Extract day, month abbreviation, and year
    day = int(nums[0])
    month = month_map[month[0].upper()]
    year = (nums[-1])
    
    # Convert the year to a 4-digit year based on the assumption
    # that the year is in the range 00-99 and represents the most recent year matching that
    if len(year) == 2:
        year = int(year)
        if year <= datetime.now().year % 100:
            year += 2000
        else:
            year += 1900
    else:
        year = int(year)
    
    # Create a datetime object
    dt = datetime(year, month, day)
     
    return dt

def parse_method_code(doc_name, study_id):
    end = doc_name.find(study_id) + len(study_id)
    method_search = re.search(r'(?![A-Za-z]*R\d)[A-Za-z]+(?:-[A-Za-z]+)*', doc_name[end:])
    if method_search is not None:
        return method_search[0]
    else:
        return None
    
def parse_fn_end(tt):
    if len(tt) > 1:
        if 'R' in tt[1]:
            rep_num = re.findall(r"\d+(?:\.\d+)?", tt[1])
            if len(rep_num):
                v = float(rep_num[0])
            else:
                v = 0
        else:
            v = 0

        if re.search(r"CO(?:[^a-zA-Z]|$)", tt[1]):
            change_order = True
        else:
            change_order = False
    else:
        v = 0
        change_order = False
    
    return change_order, v

def find_string_instance(doc, target, return_first = True):

    # function to string match - might want to improve to use regular expression or define a 'loose' matching code
    all_inds = []
    for ind, p in enumerate(doc.paragraphs):
        if target.lower() in p.text.lower():
            if return_first:
                return ind
            else:
                all_inds.append(ind)
    if return_first and len(all_inds) == 0:
        return None
    else:
        return all_inds

def table_of_contents(doc, toc_name = "table of contents"):
    toc_ind = find_string_instance(doc, toc_name)
    if toc_ind is None:
        return None, None, None

    for ind, p in enumerate(doc.paragraphs[toc_ind+1:]):
        if p.contains_page_break:
            break
    
    if ind > 40:
        inds, toc = zip(*[(ind, p.text) for ind, p in enumerate(doc.paragraphs[toc_ind:]) if bool(re.search(r'\t+\d', p.text))])
        toc_end = inds[-1] + toc_ind
        toc = list(toc)
    else:
        toc_end = toc_ind + ind - 1
        toc = [p.text for p in doc.paragraphs[toc_ind:toc_end]]

        ## issues: if format of the table does not have number tab title tab number - 

    return toc, toc_ind, toc_end

def main_paragraphs(doc, data_ind):
    find_terms = [k for k in data_ind.keys() if ('terms' in k) and ('conditions' in k)]
    if len(find_terms):
        stop = data_ind[find_terms[0]][0]
    else:
        find_appendix = [k for k in data_ind.keys() if 'appendix' in k]
        if len(find_appendix):
            stop = data_ind[find_appendix[0]][0]
        else: 
            stop = len(doc.paragraphs)

    text = ' '.join([p.text for p in doc.paragraphs[0:stop]])
    return text

def check_if_time(txt):
    found = False

    ind = txt.find('AM')
    if (ind < 5) & (ind > -1):
        found = True
    
    ind = txt.find('PM')
    if (ind < 5) & (ind > -1):
        found = True
    
    return found

def delineated_para_list(doc, delineator = ':'):
    dat = {}
    last_key = None
    for p in doc:
        split = p.text.split(delineator, 1)
        if len(split) > 1 and not check_if_time(split[-1]):
            split = p.text.split(delineator, 1)
            last_key = split[0]
            dat[last_key] = split[-1].strip()
        elif not last_key is None:
            old_val = dat[last_key]

            if old_val != '':
                dat[last_key] = old_val + ", " + p.text.strip()
            else:
                dat[last_key] = p.text.strip()
    
    return dat

def combinatorial_search_from_db(db, text, key):

        matched= []
        search = db.select(pl.col(key), pl.selectors.starts_with('Search'))
        for row in search.iter_rows():
            found = True
            for c in row[1:]:
                if c is not None:
                    vars = c.replace(' or ', '|')
                    pattern = re.compile(rf"{vars}", re.IGNORECASE)
                    matches  = pattern.search(text)
                    found = found & (matches is not None)
            
            if found:
                matched.append(row[0])
            
        return matched

class MeliorStudy():

    def __init__(self, fn, ref):
        self.filepath = fn
        self.ref = ref
        self.scrape_all_data()    
 
    def scrape_all_data(self):
        self.get_file_metadata() # directory, doc_name, ext, last_mod, created
        self.parse_filename() # study_id, study_number, client, method, study_date, v, document_type
        try:
            self.document_TOC() # doc, data_ind, levels
        except Exception as e:
            print(f'Document toc failed: {e}')
            try:
                self.create_TOC()
            except Exception as e:
                print(f'Create toc failed: {e}')
                self.issue_date = None
                self.latest_reissue = None
                self.description = None
                self.species = None
                self.sex = None
        
        if hasattr(self, 'doc'):
            self.scrape_abbreviations() # abbreviations
            self.scrape_animal_details() # sex, species, strain_add, strain_match
            self.scrape_compounds() # compounds
            self.scrape_document_details() # client add, client , description, date of issue, reissue, document type
            self.scrape_employee_details() # people, people_add
            self.scrape_methods() # extends methods
    
    def get_file_metadata(self):

        fn = os.path.basename(self.filepath)
        self.directory = os.path.dirname(self.filepath)
        self.document_name, self.ext  = os.path.splitext(fn)
        # print(self.filepath)
        last_mod = os.path.getmtime(self.filepath)
        self.last_modified =  datetime.fromtimestamp(time.mktime(time.localtime(last_mod))) # flag
        created = os.path.getctime(self.filepath)
        self.created = datetime.fromtimestamp(time.mktime(time.localtime(created))) # flag
    
    def parse_filename(self):

        p = find_study_pattern(self.document_name)
        if p:
            self.study_id = self.document_name[p.span()[0]:p.span()[1]].upper()

            split_id = self.study_id.split('_')
            if len(split_id) >= 2:
                client_code = split_id[0]
                self.study_number = int(split_id[1])
            else:
                client_code = None
                self.study_number = None

            method_code = parse_method_code(self.document_name, self.study_id)
            if method_code is not None:
                self.methods = self.match_method_code(method_code)
            else:
                self.methods = []
            self.client = self.match_client_code(client_code) 
            date_end = p.span()[-1]

            trailing_text = self.document_name[date_end:]
            tt = re.split(r'[_\s.]+', trailing_text, maxsplit=1)
            try:
                self.study_date = parse_custom_date(split_id[2].upper())
            except:
                sd = dateparser.parse(split_id[2])
                if sd is not None:
                    self.study_date  = sd # flag
                else:
                    self.study_date = None
            
            change_order, self.version = parse_fn_end(tt)
            if change_order:
                self.document_type = 'change order'
            elif 'bizdev' in self.directory.lower():
                self.document_type = 'proposal'
            elif 'o-drive' in self.directory.lower():
                self.document_type = 'report'
            else:
                self.document_type = None
        else:

            self.study_id = None
            self.study_number = None
            self.study_date = None
            self.methods = []
            self.client = None
            self.version = 0
            date_end = 0

            if 'bizdev' in self.directory.lower():
                self.document_type = 'proposal'
            elif 'o-drive' in self.directory.lower():
                self.document_type = 'report'
            else:
                self.document_type = None

    def document_TOC(self, toc_name = "TABLE OF CONTENTS", return_levels = False):
        
        # Pull out table of contents by title and formatting
        # toc_ind = find_string_instance(doc, toc_name)
        # toc = [p.text for p in doc.paragraphs[toc_ind:] if bool(re.search(r'\t+\d', p.text))]
        doc = Document(self.filepath)
        toc, toc_ind, toc_end = table_of_contents(doc, toc_name)

        # Split each line at the tab, use the length to parse the section headers
        titles = [toc[0]]
        levels = [1]
        for line in toc:
            split_line = line.split('\t')
            if len(split_line) == 3:
                sec = split_line[0]
                nums = sec.split('.')
                titles.append(split_line[-2].strip())
                levels.append(len(nums))
            elif len(split_line) == 2:
                titles.append(split_line[0].strip())
                levels.append(1)

        # Create an index with the assumption that each section has a 'heading' with the exact title in it
        index = np.array([ind for ind, paragraph in enumerate(doc.paragraphs) if paragraph.text.strip() in titles and "Heading" in paragraph.style.name]).astype(int)
        index = np.insert(index, 0, int(toc_ind))
        # Fix issues if different number of headers are found than indices
        if len(index) != len(titles):
            found = [titles.index(doc.paragraphs[i].text.strip()) for i in index]
            titles = [titles[i] for i in found]
            levels = [levels[i] for i in found]

        sort_order = np.argsort(index)
        index = index[sort_order]
        titles = [titles[i] for i in sort_order]
        levels = [levels[i] for i in sort_order]

        # Create dictionary with section headers as keys and start,stop indices as the value
        data_ind = {}
        data_ind["title page"] = (0, index[0])

        # Start and stop are interpreted as the start/stop of the entire section - i.e. subsections are included within 
        # super sections start, stop. 
        for ind, title in enumerate(titles[:-1]):

            start = index[ind]
            next_ind = [i for i,element in enumerate(levels) if element <= levels[ind] and i > ind]
            if len(next_ind):
                end = index[next_ind][0]
            else:
                end = len(doc.paragraphs)
            data_ind[title.lower()] = (start, end)

        data_ind[titles[-1].lower()] = (index[-1], len(doc.paragraphs))

        self.doc = doc
        self.data_ind = data_ind

        levels.insert(0, 1)
        levels.insert(0, 1)
        self.levels = levels
        if return_levels:
            return self.levels

    def create_TOC(self):

        doc = Document(self.filepath)
        expected_TOC = {'title page': 1,
            'table of contents': 1,
            'introduction and background': 1,
            'background': 2,
            'abbreviations used in this proposal': 2,
            'study deliverables': 1,
            'experimental procedures': 1,
            'animal description': 2,
            'housing and feeding': 2,
            'design': 2,
            'general operational terms': 2,
            'methods': 2,
            'data analysis': 2,
            'terms and conditions': 1,
            'pricing': 2,
            'terms': 2, 
            'appendix 1': 1,
            'appendix 2': 1,
            'project manager': 2,
            'project coordinator': 2,
            'client management specialist': 2,
            'appendix 3': 1,
            'executive summary': 2,
            'melior discovery overview': 2,
            'facility': 2,
            'security, monitoring, and backup capability': 2}
        
        inds = []
        titles = []
        levels = []
        for ind, p in enumerate(doc.paragraphs):
            for k in expected_TOC.keys():
                if k == p.text.lower():
                    # print(f'Exact match: {p.text}')
                    if k in titles:

                        inds.remove(inds[titles.index(k)])
                        inds.append(ind)
                        titles.remove(k)
                        titles.append(k)
                    else:
                        inds.append(ind)
                        titles.append(k)
                elif k in p.text.lower():
                    if k not in titles:
                        # print(f'Inexact match: {p.text}')
                        titles.append(k)
                        inds.append(ind)

        inds = np.array(inds)
        levels = [expected_TOC[k] for k in titles]

        sort_order = np.argsort(inds)
        inds = inds[sort_order]
        titles = [titles[i] for i in sort_order]
        levels = [levels[i] for i in sort_order]

        # Create dictionary with section headers as keys and start,stop indices as the value
        data_ind = {}
        data_ind["title page"] = (0, inds[0])

        # Start and stop are interpreted as the start/stop of the entire section - i.e. subsections are included within 
        # super sections start, stop. 
        for ind, title in enumerate(titles[:-1]):

            start = inds[ind]
            next_ind = [i for i,element in enumerate(levels) if element <= levels[ind] and i > ind]
            if len(next_ind):
                end = inds[next_ind][0]
            else:
                end = len(doc.paragraphs)
            data_ind[title.lower()] = (start, end)

        data_ind[titles[-1].lower()] = (inds[-1], len(doc.paragraphs))

        self.doc = doc
        self.data_ind = data_ind

        levels.insert(0, 1)
        self.levels = levels
    
    def scrape_document_details(self):

        end = self.data_ind['title page'][1]
        
        title_page = [p.text.strip() for p in self.doc.paragraphs[0:end] if p.text.strip() != '']
        conf = [ind for ind, x in enumerate(title_page) if 'confidentiality' in x.lower()]
        if len(conf):
            title_page = title_page[:conf[0]]

        company_line = [x.split('Prepared for ')[-1] for x in title_page if 'Prepared for' in x]
        if len(company_line) > 0:
            c = company_line[0]
            matches = self.match_client_name(c)
        else:
            # text = main_paragraphs(doc,data_ind)
            company_line = [x.split('Sponsor')[-1].split(':')[-1].strip() for x in title_page if 'Sponsor' in x]

            if len(company_line) > 0:
                matches = self.match_client_name(company_line[0]) 
            else:
                matches = self.match_client_name(' '.join(title_page)) 
        
        if (matches is None) & len(company_line):
            self.client_add = [company_line[0]]
            if self.client is None:
                self.client = [company_line[0]]
        elif matches is not None:
            client_match = matches

            if self.client:
                if self.client != client_match:
                    warnings.warn(f"Filename client codes does not match client scraped from doc. Choosing doc: \n{self.filepath}")
                    self.client = client_match
            else:
                self.client = client_match

            self.client_add = None
        else:
            self.client_add = None


        desc = self.doc.sections[0].header.paragraphs[0].text.split('\t')
        if len(desc):
            self.description = desc[0]
        else:
            self.description = None

        
        dates = [x.split(':')[-1].strip() for x in title_page if 'Date' in x]
        if len(dates) > 0:
            d1 = dateparser.parse(dates[0])
            try:
                self.issue_date = d1 # flag
                if self.study_date is None:
                    self.study_date = self.issue_date
            except Exception as e:
                print(f'Could not parse {dates[0]} as date')
                print(e)
            
        else:
            self.issue_date = None

        if len(dates) > 1:
            d2 = dateparser.parse(dates[-1])
            try:
                self.latest_reissue = d2 # flag
            except Exception as e:
                print(f'Could not parse {dates[-1]} as date')
                print(e)
        else:
            self.latest_reissue = None

        
        # first look in the filepath to see if the word report or propsoal is there
        # to determine file type

        if any(['report' in x.lower() for x in self.filepath.lower()]):
            document_type = 'report'
        elif any(['proposal' in x.lower() for x in self.filepath.lower()]):
            document_type = 'proposal'
        elif any(['change order' in x.lower() for x in self.filepath.lower()]):
            document_type = 'change order'
        else:
            # if report/proposal not in the filepath - check the title page. Changed
            # this logic 11.14.24 because I found errors due to language from proposals
            # being included in early report drafts.  
            if any(['report' in x.lower() for x in title_page]):
                document_type = 'report'
            elif any(['proposal' in x.lower() for x in title_page]):
                document_type = 'proposal'
            elif any(['change order' in x.lower() for x in title_page]):
                document_type = 'change order'
            else:
                document_type = None
    
        
        self.document_type = document_type or self.document_type

    def scrape_methods(self):
        text = main_paragraphs(self.doc, self.data_ind)
        matched_methods = self.match_methods(text) 
        if self.methods is None:
            self.methods = []

        self.methods.extend(matched_methods)  

        self.methods = list(set(self.methods))
     
    def scrape_compounds(self):
        text = main_paragraphs(self.doc, self.data_ind)
        matched_compounds = self.match_compounds(text) 
        self.compounds = matched_compounds
    
    def scrape_animal_details(self):
        self.sex = None
        self.species = None
        self.strain = None 
        find_ad = [k for k in self.data_ind.keys() if 'animal description' in k]
        if len(find_ad) == 0:
            find_ad = [k for k in self.data_ind.keys() if 'animal' in k]

        if len(find_ad):
            start,end = self.data_ind[find_ad[0]]
            ad_dict = delineated_para_list(self.doc.paragraphs[start:end])

            r = False
            m = False

            males = True
            females = False

    
            for k in ad_dict.keys():
                if 'species' in k.lower():
                    if ('rat' in ad_dict[k].lower()):
                        r = True
                    if ('mouse' in ad_dict[k].lower()) or ('mice' in ad_dict[k].lower()):
                        m = True
                    
                    if r & m:
                        self.species = 'rat and mouse'
                    elif r:
                        self.species = 'rat'
                    elif m:
                        self.species = 'mouse'
                
                
                if 'strain' in k.lower():
                    matches = self.match_strain(ad_dict[k], self.species) 
                    if len(matches):
                        self.strain = matches
                        self.strain_add = None
                    else:
                        self.strain_add = [ad_dict[k]]
                        self.strain = ad_dict[k]
        

                if 'sex' in k.lower():

                    if 'female' in ad_dict[k].lower():
                        females = True
                        if ('and' in ad_dict[k].lower()) or ('&' in ad_dict[k]):
                            males = True
                        else:
                            males = False
                    
                    if males & females:
                        self.sex = 'both'
                    elif males:
                        self.sex = 'males'
                    elif females:
                        self.sex = 'females'
                
    def scrape_employee_details(self):
        roles = ['project manager', 'project coordinator', 'principal associate', 'client management specialist']
        people = {}
        self.people_add = []
        for r in roles:
            match_keys = [k for k in self.data_ind.keys() if r in k.lower()]
            if len(match_keys):
                k = match_keys[0]
                text = ' '.join([p.text for p in self.doc.paragraphs[self.data_ind[k][0]: self.data_ind[k][1]]])
                matches = self.match_employee(text) 
                if len(matches):
                    people[r] = matches[0]
                else:
                    
                    matches = find_persons(text)
                    if len(matches):
                        people[r] = matches[0]
                        self.people_add.extend(matches)
        
        if len(people):
            self.people = people
        else:
            self.people = None
        
        if len(self.people_add) == 0:
            self.people_add = None
    
    def scrape_abbreviations(self):
    
        find_ab = [k for k in self.data_ind.keys() if 'abbreviation' in k]
        if len(find_ab):
            start,end = self.data_ind[find_ab[0]]
            ab_dict = delineated_para_list(self.doc.paragraphs[start:end])

            self.abbreviations = ab_dict
        else:
            self.abbreviations = None
       
    def get_document_entries(self):

        doc_data  = {'study_id': self.study_id, 'document_name': self.document_name, 'document_type': self.document_type, 'ext': self.ext, 'directory': self.directory, 
                      'last_modified': self.last_modified, 'created': self.created, 'filepath': self.filepath, 'version': float(self.version)}
        


        study_data  = {'study_id': self.study_id,  'study_number': self.study_number, 'study_date': self.study_date, 
                          'client': self.client,  'species': self.species, 'sex': self.sex, 'description': self.description, 
                          'issue_date': self.issue_date, 'latest_reissue': self.latest_reissue}
        
        
        return doc_data, study_data

    def get_document_data(self):
        if type(self.methods) is not list:
            self.methods = [self.methods]
        

        if type(self.compounds) is not list:
            self.compounds = [self.compounds]
        

        if type(self.strain) is not list:
            self.strain = [self.strain]

        if type(self.people) is not dict:
            self.people = {}
        
        return self.methods, self.compounds, self.people, self.strain
    
    def get_data_adds(self):
        adds = {}
        if hasattr(self, 'client_add'):
            adds['client'] =  self.client_add
        else:
            adds['client']  = None
        
        if hasattr(self, 'people_add'):
            adds['people']  = self.people_add
        else:
            adds['people'] = None

        if hasattr(self, 'strain_add'):
            adds['strain'] = self.strain_add
        else:
            adds['strain']  = None
        return adds

    def match_methods(self, text):
        
        text = text.replace('-', ' ')
        if self.ref.load_table('methods'):
            matched  = combinatorial_search_from_db(self.ref.database['methods'], text, 'method')
            return list(set(matched))
        else:
            print(f'methods not found in {self.ref.init_dir}')
    
    def match_method_code(self, text):
        if self.ref.load_table('methods'):     
            codes = self.ref.database['methods']['method_code'].str.to_uppercase()
            if text.upper() in codes:
                return self.ref.database['methods'].filter(pl.col('method_code') == text.upper())['method'].to_list()
        else:
            print(f'methods not found in {self.ref.init_dir}')

    def match_compounds(self, text):

        if self.ref.load_table('compounds'):
            matches = []
            for c in self.ref.database['compounds']['compound'].to_list():
                if c.lower() in text.lower():
                    matches.append(c)
            
            for c in self.ref.database['compounds']['Alts'].to_list():
                if c is not None:
                    if c.lower() in text.lower():
                        # matches.append(c)
                        match = self.ref.database['compounds'].filter(pl.col('Alts') == c)['compound'].item()
                        matches.append(match)

            return list(set(matches))
        else:
            print(f'compounds not found in {self.ref.init_dir}')

    def match_strain(self, text, species):
        if self.ref.load_table('strain'):

            if species is None:
                species_db = self.ref.database['strain']
            else:
                species_db = self.ref.database['strain'].filter(pl.col('species')== species)

            matched  = combinatorial_search_from_db(species_db , text, 'strain')

            return matched
        else:
            print(f'strains not found in {self.ref.init_dir}')
    
    def match_employee(self, text):
        if self.ref.load_table('employees'):
            matches = combinatorial_search_from_db(self.ref.database['employees'], text, 'employee')
            return matches
        else:
            print(f'employees not found in {self.ref.init_dir}')
    
    def match_client_name(self, text):
        if self.ref.load_table('clients'):
        
            matches = combinatorial_search_from_db(self.ref.database['clients'], text, 'client')

            if len(matches) > 1:
                warnings.warn("More than one client found in text, returning first")
                return matches[0]
            elif len(matches) == 0:
                return None
            else:
                return matches[0]
        else:
            print(f'clients not found in {self.ref.init_dir}')
    
    def match_client_code(self, text):
        if self.ref.load_table('clients'):
        
            matches = []
            for c in self.ref.database['clients']['scrape'].to_list():
                if c.lower() == text.lower():
                    code = self.ref.database['clients'].filter(pl.col('scrape') == c)['client'].to_list()
                    matches.extend(code)
            
            for c in self.ref.database['clients']['alt 1'].to_list():
                if c is not None:
                    if c.lower() == text.lower():
                        code = self.ref.database['clients'].filter(pl.col('alt 1') == c)['client'].to_list()
                        matches.extend(code)
            
            for c in self.ref.database['clients']['alt 2'].to_list():
                if c is not None:
                    if c.lower() == text.lower():
                        code = self.ref.database['clients'].filter(pl.col('alt 2') == c)['client'].to_list()
                        matches.extend(code)


            if len(matches) > 1:
                warnings.warn("More than one client found in text, returning first")
                return matches[0]
            elif len(matches) == 0:
                return None
            else:
                return matches[0]
        else:
            print(f'clients not found in {self.ref.init_dir}')
    
class buildingAPI(CsvDatabase):

    def __init__(self, folder, ref_db = None):
        if folder is None:
            folder = "/Databases"

        if ref_db is None:
            self.ref_db = CsvDatabase("/Databases")
        else:
            self.ref_db = CsvDatabase(folder = ref_db)
            self.ref_db.load_all()

        schema = {
            
            'documents' : {'document_id': str, 'study_id': str, 'document_number': int, 'document_name':str, 'document_type':str, 'ext': str, 'directory': str, 'last_modified': datetime,
                                   'created': datetime, 'filepath': str, 'version': float},
            'studies': {'study_id': str, 'study_number': int, 'study_date': datetime, 
                        'client': str, 'species': str, 'sex': str,'description': str, 
                        'proposal_id': str, 'proposal_issue_date': datetime, 'proposal_latest_reissue': datetime, 
                        'report_id': str, 'report_issue_date': datetime, 'report_latest_reissue': datetime},
            'study_employees': {'study_id' : str, 'employee': str, 'role': str},
            'study_methods': {'study_id':str, 'method':str},
            'study_compounds' : {'study_id':str, 'compound': str},
            'study_strains': {'study_id': str, 'strain': str},
            'scraped_files': {'filepath': str, 'success': bool}

        }

        super().__init__(folder, schema = schema)
        self.load_database()

    def load_database(self):
        for tbl in self.schema.keys():
            if tbl not in self.database.keys():
                self.create_tbl(tbl, data = None, schema = self.schema[tbl])

    def get_document_number(self, ms):

        if self.load_table('documents'):
            if len(self.database['documents']) != 0:
                filter_by = {'study_id': ms.study_id, 'document_type': ms.document_type}
                found, target_docs = self.get_entries('documents', filter_by, as_dict = False)
                last_doc_num = target_docs.select('document_number')
                if len(last_doc_num):
                    return last_doc_num.max().item() + 1
                else:
                    return 0
            else:
                return 0
            
    def is_latest_document(self, ms):

        if self.load_table('documents'):
            if len(self.database['documents']) != 0:
                filter_by = {'study_id': ms.study_id, 'document_type': ms.document_type}
                found, target_docs = self.get_entries('documents', key =filter_by, as_dict = False)
                if len(target_docs) != 0:
                    if ms.version:
                        if target_docs.select('version').max().item() < ms.version:
                            return True
                        elif target_docs.select('version').max().item() > ms.version:
                            return False
                
                    # This line MUST be less than or equal to because 
                    if ms.last_modified >= target_docs.select('last_modified').max().item():
                        return True
                    else:
                        return False
                else:
                    return True
            else:
                return True
        else:
            print(f'filedata not found in {self.init_dir}')
    
    def infer_study_id(self, ms):

        month_map = {
            1:'JAN', 2: 'FEB',3: 'MAR', 4: 'APR', 
            5: 'MAY',6: 'JUN',7: 'JUL',8: 'AUG',
            9 : 'SEPT', 10: 'OCT', 11: 'NOV', 12: 'DEC'
        }
        if ms.client:
            filter_client = {'client': ms.client}
            found, matches = self.ref_db.get_entries('clients', filter_client)
            if len(matches):
                code = matches.select(pl.col('client_code')).item()
            else:
                code = 'UNKNOWN'
        else:
            code = 'UNKNOWN'
        
        if ms.study_date:

            date = ms.study_date.strftime('%d') + month_map[ms.study_date.month] + ms.study_date.strftime('%y')
        else:
            date = ms.created.strftime('%d') + month_map[ms.created.month] + ms.created.strftime('%y')
        
        if len(self.database['studies']):
            all_ids = self.database['studies'].select(pl.col('study_id').str.split('_').list.to_struct()).unnest('study_id')
            all_ids = all_ids.with_columns([(pl.col("field_2").map_elements(lambda x: dateparser.parse(x), return_dtype = datetime).alias('date'))]).unique()

            match = all_ids.filter(pl.col('field_0') == code)
       
            if match.is_empty():
                study_num = '00'
            if not match.filter(pl.col("date").dt.date() == ms.study_date.date()).is_empty():
                study_num = match.filter(pl.col("date").dt.date() == ms.study_date.date()).select('field_1').item()
                date = match.filter(pl.col("date").dt.date() == ms.study_date.date()).select('field_2').item()
            else:
                study_num = str(int(match.select('field_1').max().item())+1).zfill(2)
        else:
            study_num = '00'
        
        return code+ '_'+study_num+ '_' + date, study_num, date 
        
    # Finding if specific elements are in the database
    def has_filepath(self, fp):
        if self.load_table('documents'):
            found, dn = self.get_entries('documents', {'filepath': fp})
            if found:
                if len(dn) == 1:
                    return dn['document_number'].item()
                elif len(dn) == 0:
                    return None
                else:
                    print('Multiple document numbers associated with single filepath')
                    return None
                    # raise DatabaseError(f'Multiple document numbers associated with single filepath: {fp}')
                    # return dn['document_number'].to_list()
            else:
                return None
        else:
            print(f'filedata not found in {self.init_dir}')
            return None
        
    def has_study(self, study_id):
        if self.load_table('studies'):
        
            if study_id in self.database['studies']['study_id']:
                return True
            else:
                return False
        
        else:
            print(f'studies not found in {self.init_dir}')

    # adding document 
    def add_filepath(self, filepath, rescrape = False):

        # initialize booleans indicating if data was created or updated
        updated_data = False
        created_data = False

        # create a MeliorStudy object using filepath and reference
        ms = MeliorStudy(filepath, self.ref_db)

        # check to see if the code was able to get anything beyond details from filename
        scraped = hasattr(ms, 'doc')

        # determine the study ID if Melior Study didn't successfully scrape it
        if ms.study_id is None: # if study_id not part of file, infer it
            ms.study_id, ms.study_number, study_date = self.infer_study_id(ms)
            if ms.study_date is None:
                ms.study_date = study_date
        
        # new is boolean indicating if a ms is a new entry. updates contains dict of all new things added or entries updated in documents table
        # doc_id is the document id in documents table
        new, updates, doc_id = self.update_documents(ms, update = rescrape) 
        print(f'document updates: {updates}')

        # if nothing new was created and we're not rescraping, then return study id, and booleans to indicate if it was created or updated
        if (not new) and (not rescrape):
            return ms.study_id, updated_data, created_data, None
        else: 
            if scraped:
                # if the study id exists and this document is the latest document, update the study information. 
                # if study id exists and this document is not the latests, do nothing.
                if self.has_study(ms.study_id): 
                    if self.is_latest_document(ms):
                        self.update_study_with_document(ms, doc_id, delete_old = rescrape) 
                        updated_data = True
                # if study id not in databse, create a new study using this document
                else: 
                    created_data = self.create_study_with_document(ms, doc_id)

        return ms.study_id, updated_data, created_data, ms.get_data_adds()

    def create_study_with_document(self, ms, doc_id):

        doc_data, study_data = ms.get_document_entries()
        if doc_data['document_type'] == 'proposal':
            study_data['proposal_id'] = doc_id
            study_data['proposal_issue_date'] = study_data.pop('issue_date')
            study_data['proposal_latest_reissue'] = study_data.pop('latest_reissue')
        elif doc_data['document_type'] == 'report':
            study_data['report_id'] = doc_id
            study_data['report_issue_date'] = study_data.pop('issue_date')
            study_data['report_latest_reissue'] = study_data.pop('latest_reissue')
        else:
            print(f'{study_data["document_name"]} is not a report or proposal')
            return False
        
        self.write_entry('studies', key = {'study_id': ms.study_id}, data = study_data)

        methods, compounds, people, strain = ms.get_document_data()
        for m in methods:
            if m is not None:
                self.write_entry('study_methods', key = {'study_id': ms.study_id, 'method': m}, data = {'study_id': ms.study_id, 'method': m})
        for c in compounds:
            if c is not None:
                self.write_entry('study_compounds', key = {'study_id': ms.study_id, 'compound': c}, data = {'study_id': ms.study_id, 'compound': c})
        for s in strain:
            if s is not None:
                self.write_entry('study_strains', key = {'study_id': ms.study_id, 'strain': s}, data ={'study_id': ms.study_id, 'strain': s})
        for k,v in people.items():
            if v is not None:
                if k is not None:
                    self.write_entry('study_employees', key = {'study_id': ms.study_id, 'role': k}, data = {'study_id': ms.study_id, 'employee': v, 'role': k})

        return True
    
    def update_study_with_document(self, ms, doc_id, delete_old = True):

        doc_data, study_data = ms.get_document_entries()
        if doc_data['document_type'] == 'proposal':
            study_data['proposal_id'] = doc_id
            study_data['proposal_issue_date'] = study_data.pop('issue_date')
            study_data['proposal_latest_reissue'] = study_data.pop('latest_reissue')
        elif doc_data['document_type'] == 'report':
            study_data['report_id'] = doc_id
            study_data['report_issue_date'] = study_data.pop('issue_date')
            study_data['report_latest_reissue'] = study_data.pop('latest_reissue')
        else:
            print(f'{study_data["document_name"]} is not a report or proposal')
            return False, (None, None, None)
        
        # study_data = {k:v for (k,v) in study_data.items() if v is not None}
        updates = self.update_entry('studies', key = {'study_id': ms.study_id}, data = study_data)
        print(f'study updates: {updates}')
        

        # if (doc_data['document_type'] == 'report') and (delete_old == False):
        #     delete_old = True
        methods, compounds, people, strain = ms.get_document_data()
        self.reconcile_list('study_methods', study_data['study_id'], methods, 'method', delete_old = delete_old) 
        self.reconcile_list('study_compounds', study_data['study_id'], compounds, 'compound', delete_old = delete_old)
        self.reconcile_list('study_strains', study_data['study_id'], strain, 'strain', delete_old = delete_old)

        for k, v in people.items():
            
            if k is not None:
                found, db_data = self.get_entries('study_employees', {'study_id': ms.study_id, 'role': k}, as_dict = True)
                if found & (db_data['employee'] != v) & (v is not None):
                    # self.delete_entries('study_employees', {'study_id': ms.study_id, 'role': k})
                    self.write_entry('study_employees', key = {'study_id': ms.study_id, 'role': k}, data = {'study_id': ms.study_id, 'role': k, 'employee': v}, overwrite = True)
                    
        
        return True

    def update_documents(self, ms, update = True):

        # load the data affilicated with the MeliorStudy
        doc_data, _ = ms.get_document_entries()

        # determine if document is already in the database. doc_num == None means its not
        doc_num = self.has_filepath(ms.filepath)

        # if document is not in the database, determine its number and document type to generate a unique document ID
        # then add the data to 'documents' table using 'filepath' as the key
        if doc_num is None:
            doc_num = self.get_document_number(ms)
            doc_data['document_number']  = doc_num
            if doc_data['document_type'] is not None:
                let = doc_data['document_type'][0].upper()
            else:
                let = 'U'
            doc_data['document_id'] = doc_data['study_id']+ '_' + let + str(doc_num).zfill(2)
            written =  self.write_entry('documents',  data = doc_data,  key = {'filepath': doc_data['filepath']}, overwrite = False)
            if written:
                updates = doc_data
            else:
                warnings.warn('Could not write document')
                return False, None, doc_num
        # If document was in the database, we might want to update the database. Written will be set to false.
        else:
            written = False
            found, db_data = self.get_entries('documents', {'filepath': doc_data['filepath']}, as_dict = True)
            # if the filepath has been modified since its record in the database or if are rescraping with the update = True,
            # then update the documents table
            if (doc_data['last_modified']  > db_data['last_modified']) or update:
                doc_data['document_number'] = doc_num
                updates = self.update_entry('documents', key = {'filepath': doc_data['filepath']}, data = doc_data)
            else:
                print(f'{doc_data["document_name"]} already in database')
                updates = None
        
        # written is a boolean if a new entry was written, updates containds doc_data for new entries and dict of changes for updates. 
        return written, updates, doc_data['document_id']
 
    def reconcile_list(self, tbl, study_id, dat, dat_key, delete_old = False):
  
        if type(dat) != list:
            dat = [dat]

        found, db = self.get_entries(tbl, {'study_id': study_id}, as_dict= True)
        db_data = db[dat_key]

        if type(db_data) != list:
            db_data = [db_data]

        new_items = [x for x in dat if x not in db_data]
        for x in new_items:
            if x is not None:
                self.write_entry(tbl,  data = {'study_id': study_id, dat_key: x})
        
        if delete_old:
            unmatched_items = [x for x in db_data if x not in dat]
            # print(unmatched_items)
            deleted  = self.delete_entries(tbl, {'study_id': study_id, dat_key: unmatched_items})
            if len(deleted):
                print(f'Deleted: {deleted[dat_key].to_list()}')
            # self.database[tbl] = self.database[tbl].filter(~((self.database[tbl]['study_id']== study_id) & (self.database[tbl][dat_key].is_in(unmatched_items))))

    # Bulk adding of data by searching folder. Will search for all .docx in dirpath
    # if rescrape is False, only filepaths NOT previously scraped will be examined
    # if rescrape is True, ALL filepaths will be scraped
    def scrape_folder(self, dirpath, rescrape = False): 
        add_log = []
        for filename in glob.iglob(os.path.join(dirpath, '**', '*.docx'), recursive=True):
            if '~$' not in filename:

                attempt = True
                found, db_data = self.get_entries('scraped_files', {'filepath':filename}, as_dict= False)
                if found:
                    if len(db_data):
                        if rescrape is False:
                            attempt = False

                if attempt:
                    try:
                        # Updated_data: boolen if a study was updated using this document. (false if document is latest for study)
                        # created_data: boolean if a study was created using this document.
                        # adds: dict of things found in document that aren't present in the references
                        study_id, updated_data, created_data, adds = self.add_filepath(filename, rescrape = rescrape)
                        self.write_entry('scraped_files', key = {'filepath': filename}, data = {'filepath': filename, 'success' : True}, overwrite = rescrape)

                        for k,v in adds.items():
                            if v is not None:
                                for v2 in v:
                                    add_log.append({'study_id': study_id,  'add_type': k, 'add': v2, 'filepath': filename,})
                        
                    except Exception as e:
                        print(f'Exception occurred when adding {filename}: {e}')
                        self.write_entry('scraped_files', key = {'filepath': filename}, data = {'filepath': filename, 'success' : False})
    
        
        return pl.DataFrame(add_log)


