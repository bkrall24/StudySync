
from docxtpl import DocxTemplate, subdoc, InlineImage
from docx import Document, table, text
from docx.shared import Inches, Mm
import polars as pl
import os

def get_matching_bio(name, template, bios = '/Volumes/Company/Becca/Study Database/Proposal_Bios.docx'):
    bio_doc = Document(bios)
    starts = [ind for ind, p in enumerate(bio_doc.paragraphs) if name.lower() in p.text.lower()]
    
    
    if len(starts):
        ends = [ind for ind, p in enumerate(bio_doc.paragraphs) if 'BIOEND' in p.text and ind > starts[0]]
        if len(ends):
            subdoc = template.new_subdoc()
            for p in bio_doc.paragraphs[starts[0]+1: ends[0]]:
                    subdoc.add_paragraph(p.text, p.style)
        
            return subdoc
        else:
            return None
    else:
         return None

def get_matching_pic(name, pics = '/Volumes/Company/Becca/Study Database/Bio_Pics',
                     refer ="/Volumes/Company/Becca/Study Database/Bio_Pics/Bio_Pic_Dimensions.csv"):
    reference = pl.read_csv(refer)
    filenames = os.listdir(pics)
    possible_pics =  [os.path.splitext(x)[0] for x in filenames]
    if name in possible_pics:
        fn = filenames[possible_pics.index(name)]
        fp = os.path.join(pics, fn)
        
        matched = reference.filter(pl.col('Name') == fn)
        if len(matched):
            ratio = matched['x'].item() / matched['y'].item()
            return fp, ratio
        else:
            return fp, None
    else:
        #  print(f"{name} not found")
         return None, None

def fill_proposal_template(data_dict, save_path = "/Users/rebeccakrall/Desktop/test_filling.docx", 
                           template_path = "/Volumes/Company/Becca/Study Database/Proposal_template_update_dxpt.docx"):
    template = DocxTemplate(template_path)
    place = template.get_undeclared_template_variables()
    replace = {k : None for k in place}

    for k in replace.keys():
        if k in data_dict.keys():
             replace[k] = data_dict[k]

    height = 60
    if data_dict['pm'] is not None:
        bio = get_matching_bio(data_dict['pm'], template = template)
        if bio is not None:
             replace['pm_bio'] = bio
        pic, aspect = get_matching_pic(data_dict['pm'])
        if pic is not None:
        #      replace['pm_image']  = pic
            if aspect:
                width = aspect * height
            else:
                height = 50
            replace['pm_image']  = InlineImage(template, image_descriptor=pic, width=Mm(width), height=Mm(height))

    if data_dict['pc'] is not None:
        bio = get_matching_bio(data_dict['pc'], template = template)
        if bio is not None:
             replace['pc_bio'] = bio
        pic, aspect = get_matching_pic(data_dict['pc'])        
        if pic is not None:
            if aspect:
                width = aspect * height
            else:
                height = 50
            replace['pc_image']  = InlineImage(template, image_descriptor=pic, width=Mm(width), height=Mm(height))

    if data_dict['cms'] is not None:
        bio = get_matching_bio(data_dict['cms'], template = template)
        if bio is not None:
             replace['cms_bio'] = bio
        pic, aspect = get_matching_pic(data_dict['cms'])
        if pic is not None:
            if aspect:
                width = aspect * height
            else:
                height = 50
            replace['cms_image']  = InlineImage(template, image_descriptor=pic, width=Mm(width), height=Mm(height))
    

    template.render(replace)
    template.save(save_path)